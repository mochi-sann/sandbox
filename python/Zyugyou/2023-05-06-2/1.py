import docx
doc = docx.Document('ch13/demo.docx')
print(len(doc.paragraphs))
for para in doc.paragraphs:
    print(para.text)
